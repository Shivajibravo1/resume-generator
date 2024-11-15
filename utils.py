import streamlit as st
import base64
from io import BytesIO
from fpdf import FPDF, HTMLMixin
from docx import Document
from bs4 import BeautifulSoup

def apply_custom_css():
    st.markdown("""
    <style>
    .stButton > button {
        background-color: #3498db;
        color: white;
        font-weight: bold;
        border-radius: 5px;
        border: none;
        padding: 10px 20px;
        transition: background-color 0.3s ease;
    }
    .stButton > button:hover {
        background-color: #2980b9;
    }
    .stTextInput > div > div > input {
        background-color: #f0f2f6;
        border-radius: 5px;
        border: 1px solid #bdc3c7;
    }
    .stTextArea > div > div > textarea {
        background-color: #f0f2f6;
        border-radius: 5px;
        border: 1px solid #bdc3c7;
    }
    .stDateInput > div > div > input {
        background-color: #f0f2f6;
        border-radius: 5px;
        border: 1px solid #bdc3c7;
    }
    </style>
    """, unsafe_allow_html=True)

class PDF(FPDF, HTMLMixin):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Professional Resume', 0, 1, 'C')

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

def create_download_link(content, file_format):
    if file_format == "PDF":
        pdf = PDF()
        pdf.add_page()
        pdf.write_html(content)
        
        buffer = BytesIO()
        pdf.output(buffer)
        b64 = base64.b64encode(buffer.getvalue()).decode()
        return f'<a href="data:application/pdf;base64,{b64}" download="resume.pdf" class="download-button">Download PDF</a>'
    elif file_format == "DOCX":
        doc = Document()
        soup = BeautifulSoup(content, 'html.parser')
        for element in soup.find_all(['h1', 'h2', 'p', 'ul']):
            if element.name == 'h1':
                doc.add_heading(element.text, level=1)
            elif element.name == 'h2':
                doc.add_heading(element.text, level=2)
            elif element.name == 'p':
                doc.add_paragraph(element.text)
            elif element.name == 'ul':
                for li in element.find_all('li'):
                    doc.add_paragraph(li.text, style='List Bullet')
        
        buffer = BytesIO()
        doc.save(buffer)
        b64 = base64.b64encode(buffer.getvalue()).decode()
        return f'<a href="data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}" download="resume.docx" class="download-button">Download DOCX</a>'